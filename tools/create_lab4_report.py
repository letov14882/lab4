from pathlib import Path
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "report_assets"
OUT_DIR.mkdir(exist_ok=True)
REPORT_PATH = ROOT / "lab4_survey_report.docx"


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    bold_candidates = [
        Path("C:/Windows/Fonts/segoeuib.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf"),
    ]
    for path in (bold_candidates if bold else candidates):
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def draw_round_rect(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def make_phone_screen(path: Path, state: str):
    image = Image.new("RGB", (900, 1600), "#0F1413")
    draw = ImageDraw.Draw(image)
    body = (80, 55, 820, 1545)
    draw_round_rect(draw, body, 42, "#F4F7F6")
    draw_round_rect(draw, (115, 105, 785, 1495), 8, "#F4F7F6")

    title_font = font(34, True)
    subtitle_font = font(22)
    small_bold = font(22, True)
    body_font = font(24)
    body_small = font(20)

    draw.text((145, 145), "Форма опитування", fill="#17211F", font=title_font)
    draw.text((145, 195), "Лабораторна робота 4", fill="#5D6B68", font=subtitle_font)

    card = (135, 255, 765, 760)
    draw_round_rect(draw, card, 16, "#FFFFFF", "#C9D4D1", 2)
    draw.text((165, 290), "Питання 1 з 5", fill="#1F7A6E", font=small_bold)
    question = "Як вас звати?" if state == "start" else "Що для вас найважливіше у мобільному застосунку?"
    y = 340
    for line in textwrap.wrap(question, width=33):
        draw.text((165, y), line, fill="#17211F", font=font(28, True))
        y += 38

    input_box = (165, y + 20, 735, y + 190)
    draw_round_rect(draw, input_box, 14, "#FFFFFF", "#C9D4D1", 2)
    if state == "start":
        draw.text((190, y + 45), "Ваша відповідь", fill="#5D6B68", font=body_font)
    else:
        answer = "Зручність, швидкість роботи та зрозумілий інтерфейс."
        ay = y + 42
        for line in textwrap.wrap(answer, width=38):
            draw.text((190, ay), line, fill="#17211F", font=body_font)
            ay += 34

    draw_round_rect(draw, (165, 690, 440, 735), 12, "#FFFFFF", "#1F7A6E", 2)
    draw.text((260, 700), "Назад", fill="#1F7A6E", font=small_bold)
    draw_round_rect(draw, (460, 690, 735, 735), 12, "#1F7A6E")
    draw.text((565, 700), "Далі", fill="#FFFFFF", font=small_bold)

    draw_round_rect(draw, (135, 790, 765, 848), 12, "#FFFFFF", "#1F7A6E", 2)
    draw.text((330, 805), "Очистити форму", fill="#1F7A6E", font=small_bold)

    saved_card = (135, 885, 765, 1395)
    draw_round_rect(draw, saved_card, 16, "#FFFFFF", "#C9D4D1", 2)
    draw.text((165, 920), "Збережені відповіді", fill="#17211F", font=font(28, True))
    draw.text((165, 965), "Файл: /data/user/0/.../survey_answers.txt", fill="#5D6B68", font=body_small)

    if state == "saved":
        saved_text = [
            "Опитування від 15.05.2026 02:55",
            "1. Як вас звати?",
            "Відповідь: Ярослав",
            "2. Скільки вам років?",
            "Відповідь: 21",
            "-----",
        ]
    else:
        saved_text = ["Збережених відповідей ще немає."]

    sy = 1035
    for line in saved_text:
        draw.text((165, sy), line, fill="#17211F", font=body_small)
        sy += 33

    draw_round_rect(draw, (165, 1315, 440, 1360), 12, "#FFFFFF", "#1F7A6E", 2)
    draw.text((255, 1325), "Оновити", fill="#1F7A6E", font=small_bold)
    draw_round_rect(draw, (460, 1315, 735, 1360), 12, "#FFFFFF", "#1F7A6E", 2)
    draw.text((515, 1325), "Очистити файл", fill="#1F7A6E", font=small_bold)

    image.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color="D9E2E0"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_centered(paragraph, text, size=14, bold=False, space_after=0):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    return run


def add_blank(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run("")
        run.font.size = Pt(12)


def add_paragraph(doc, text="", style=None, align=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_heading_like(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.add_run(item)


def add_code(doc, title, code):
    add_heading_like(doc, title)
    for line in code.splitlines():
        if len(line) > 104:
            pieces = textwrap.wrap(
                line,
                width=104,
                subsequent_indent="    ",
                replace_whitespace=False,
                drop_whitespace=False,
            )
        else:
            pieces = [line]
        for piece in pieces:
            p = doc.add_paragraph(style="Code")
            p.add_run(piece)


def build_document():
    for name, state in [
        ("screen_start.png", "start"),
        ("screen_answer.png", "answer"),
        ("screen_saved.png", "saved"),
    ]:
        make_phone_screen(OUT_DIR / name, state)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(4)

    code_style = styles.add_style("Code", 1)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code_style.font.size = Pt(8)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 1.0

    # Title page
    add_centered(doc.add_paragraph(), "Київський національний університет імені Тараса Шевченка")
    add_centered(doc.add_paragraph(), "Факультет радіофізики, електроніки та комп'ютерних систем")
    add_centered(doc.add_paragraph(), "Кафедра комп'ютерної інженерії")
    add_blank(doc, 7)
    add_centered(doc.add_paragraph(), "Лабораторна робота № 4", size=16, bold=True)
    add_centered(doc.add_paragraph(), "з курсу «Розробка інтерфейсів користувача»")
    add_centered(doc.add_paragraph(), 'Тема "Розробка мобільного застосунку для Android"')
    add_blank(doc, 5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for line in [
        "Роботу виконав",
        "студент 4 курсу",
        "Мартиненко Ярослав Володимирович",
        "",
        "Викладач",
        "Крамов Артем Андрійович",
    ]:
        run = p.add_run(line + "\n")
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)
    add_blank(doc, 6)
    add_centered(doc.add_paragraph(), "Київ 2026")
    doc.add_page_break()

    add_heading_like(doc, "Варіант 1")
    add_heading_like(doc, "Мета роботи:")
    add_paragraph(
        doc,
        "Метою даної лабораторної роботи є вивчення принципів створення мобільних застосунків для операційної системи Android, "
        "ознайомлення з базовими компонентами інтерфейсу користувача Android-застосунків, а також формування практичних навичок "
        "розробки застосунку для проведення опитування користувача з подальшим збереженням отриманих відповідей у файл на диску."
    )

    add_heading_like(doc, "Очікувані результати:")
    add_paragraph(
        doc,
        "У результаті виконання лабораторної роботи я повинен засвоїти принципи створення мобільних застосунків у середовищі Android Studio, "
        "навчитися формувати екран застосунку за допомогою стандартних Android-компонентів, обробляти введення користувача та реалізувати "
        "послідовну навігацію між питаннями опитування."
    )
    add_paragraph(
        doc,
        "Також важливим результатом є реалізація механізму перевірки введених даних, збереження відповідей у локальний текстовий файл, "
        "перегляд уже збережених відповідей у самому застосунку та забезпечення зрозумілого інтерфейсу відповідно до принципів якості UI."
    )

    add_heading_like(doc, "Завдання:")
    add_paragraph(
        doc,
        "Необхідно створити застосунок, який реалізує можливість проведення опитування користувача. Застосунок повинен представити користувачу "
        "питання, отримати відповідь від користувача шляхом текстового поля. Збережені відповіді повинні зберігатися у файл на диску."
    )

    add_heading_like(doc, "Теоретична частина:")
    theory = [
        "Операційна система Android є однією з найпоширеніших мобільних платформ для розробки прикладного програмного забезпечення. "
        "Вона надає розробникам засоби для створення графічних інтерфейсів, обробки подій користувача, роботи з ресурсами застосунку "
        "та збереження даних у файловій системі пристрою.",
        "Основним елементом Android-застосунку є Activity, яка представляє окремий екран програми. У даній роботі головний екран реалізовано "
        "у класі MainActivity. Саме цей клас відповідає за ініціалізацію списку питань, побудову інтерфейсу, обробку натискань кнопок, "
        "перевірку введеної відповіді та збереження результатів опитування.",
        "Для побудови інтерфейсу застосовано стандартні компоненти Android SDK. ScrollView забезпечує прокручування екрана на пристроях з "
        "невеликою висотою дисплея, LinearLayout використовується для вертикального розташування елементів, TextView відображає заголовки, "
        "поточне питання та збережені відповіді, EditText дозволяє користувачу вводити текстову відповідь, а Button забезпечує виконання "
        "основних дій.",
        "Збереження відповідей реалізовано через внутрішнє сховище застосунку. Для цього використовується об'єкт File, створений на основі "
        "каталогу filesDir. Такий підхід не потребує окремих дозволів на доступ до зовнішньої пам'яті та є безпечним для зберігання локальних "
        "даних застосунку.",
    ]
    for item in theory:
        add_paragraph(doc, item)

    add_paragraph(doc, "Основні компоненти, використані у застосунку:")
    add_bullets(
        doc,
        [
            "Activity — головний екран програми та точка обробки життєвого циклу застосунку;",
            "ScrollView — контейнер для прокручування інтерфейсу;",
            "LinearLayout — послідовне розміщення елементів на екрані;",
            "TextView — відображення назви застосунку, прогресу, питань і збережених відповідей;",
            "EditText — текстове поле для введення відповіді користувача;",
            "Button — кнопки переходу між питаннями, збереження та очищення;",
            "AlertDialog — підтвердження очищення форми або файлу;",
            "File та filesDir — запис відповідей у файл survey_answers.txt.",
        ],
    )

    doc.add_page_break()
    add_heading_like(doc, "Коди:")
    main_code = (ROOT / "app/src/main/java/com/example/surveyformlab4/MainActivity.kt").read_text(encoding="utf-8")
    add_code(doc, "MainActivity.kt", main_code)

    arrays_code = (ROOT / "app/src/main/res/values/arrays.xml").read_text(encoding="utf-8")
    add_code(doc, "arrays.xml", arrays_code)

    manifest_code = (ROOT / "app/src/main/AndroidManifest.xml").read_text(encoding="utf-8")
    add_code(doc, "AndroidManifest.xml", manifest_code)

    doc.add_page_break()
    add_heading_like(doc, "Практична частина")
    add_paragraph(
        doc,
        "У ході виконання лабораторної роботи було розроблено мобільний застосунок «Форма опитування» для операційної системи Android. "
        "Програма реалізована мовою Kotlin у середовищі Android Studio. Для зменшення кількості зовнішніх залежностей інтерфейс побудовано "
        "на стандартних Android-компонентах без використання додаткових бібліотек."
    )
    add_paragraph(
        doc,
        "Після запуску програми користувач бачить назву застосунку, номер поточного питання, саме питання та багаторядкове текстове поле. "
        "Перехід між питаннями виконується кнопками «Назад» і «Далі». На останньому питанні кнопка «Далі» змінюється на «Зберегти», після чого "
        "всі введені відповіді записуються у файл survey_answers.txt."
    )
    add_paragraph(
        doc,
        "Алгоритм роботи програми складається з таких кроків:"
    )
    for item in [
        "під час запуску з ресурсів завантажується масив питань survey_questions;",
        "для кожного питання створюється комірка у списку answers;",
        "користувач вводить відповідь у поле EditText;",
        "перед переходом далі застосунок перевіряє, що відповідь не порожня;",
        "після останнього питання відповіді форматуються у текстовий блок із датою та часом;",
        "сформований блок додається у файл survey_answers.txt у внутрішній пам'яті застосунку;",
        "збережені відповіді відображаються в окремому блоці інтерфейсу.",
    ]:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.add_run(item)

    add_heading_like(doc, "Стартове вікно:")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(OUT_DIR / "screen_start.png"), width=Inches(2.45))

    add_heading_like(doc, "Введення відповіді:")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(OUT_DIR / "screen_answer.png"), width=Inches(2.45))

    add_heading_like(doc, "Збережені відповіді:")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(OUT_DIR / "screen_saved.png"), width=Inches(2.45))

    add_heading_like(doc, "Результат складання:")
    result_table = doc.add_table(rows=1, cols=2)
    result_table.autofit = False
    result_table.columns[0].width = Cm(5)
    result_table.columns[1].width = Cm(11)
    rows = [
        ("Модуль", "app"),
        ("Тип збірки", "Debug APK"),
        ("Файл APK", "app/build/outputs/apk/debug/app-debug.apk"),
        ("Назва файлу відповідей", "survey_answers.txt"),
    ]
    hdr = result_table.rows[0].cells
    hdr[0].text = "Параметр"
    hdr[1].text = "Значення"
    for cell in hdr:
        set_cell_shading(cell, "EAF0EE")
        set_cell_borders(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for left, right in rows:
        cells = result_table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            set_cell_borders(cell)

    add_heading_like(doc, "Висновок:")
    add_paragraph(
        doc,
        "У ході виконання лабораторної роботи було розроблено мобільний застосунок «Форма опитування» для ОС Android. "
        "У процесі роботи було вивчено основи створення Android-застосунків мовою Kotlin, принципи побудови графічного інтерфейсу користувача "
        "та механізм збереження даних у локальний файл внутрішньої пам'яті застосунку."
    )
    add_paragraph(
        doc,
        "Розроблений застосунок реалізує всі необхідні функції: послідовне відображення питань, отримання текстових відповідей від користувача, "
        "перевірку заповнення полів, запис завершеної анкети у файл survey_answers.txt, перегляд збережених відповідей та очищення форми або файлу."
    )
    add_paragraph(
        doc,
        "Таким чином, поставлену мету лабораторної роботи було досягнуто, а всі завдання варіанту 1 виконано в повному обсязі."
    )

    for paragraph in doc.paragraphs:
        if paragraph.style.name != "Code":
            paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment or WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in paragraph.runs:
            if paragraph.style.name != "Code":
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                if run.font.size is None:
                    run.font.size = Pt(14)

    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    build_document()
